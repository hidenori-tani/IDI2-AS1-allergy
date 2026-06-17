"""Build Cover_Letter_AllergolInt.docx and Tables_AllergolInt.docx
for the Allergology International submission, matching the format of
`Cover_Letter_TANI.docx` (Times New Roman / 11 pt / single paragraph spacing).
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
OUT_COVER = ROOT / "Cover_Letter_AllergolInt.docx"
OUT_TABLES = ROOT / "Tables_AllergolInt.docx"
OUT_COVER_BBRC = ROOT / "Cover_Letter_BBRC.docx"
OUT_TABLES_BBRC = ROOT / "Tables_BBRC.docx"
OUT_HIGHLIGHTS_BBRC = ROOT / "Highlights_BBRC.docx"
OUT_DOI_BBRC = ROOT / "Declaration_of_Interest_BBRC.docx"


def set_default_font(doc: Document, name: str = "Times New Roman", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    # east-asia / complex-script fallback
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def add_para(doc: Document, text: str = "", bold: bool = False, italic: bool = False,
             align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


# --------------------------------------------------------------------------
# 1) Cover letter
# --------------------------------------------------------------------------
def build_cover_letter() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Header: addressee
    add_para(doc, "The Editor-in-Chief")
    add_para(doc, "Allergology International")
    add_para(doc, "Japanese Society of Allergology")
    add_para(doc, "")

    # Date
    add_para(doc, "20 April, 2026")
    add_para(doc, "")

    # Salutation
    add_para(doc, "Dear Editor-in-Chief:")
    add_para(doc, "")

    # Body paragraphs
    add_para(
        doc,
        "Please find enclosed my manuscript, \u201cBulk tissue transcriptomes obscure the "
        "IDI2-AS1 / IL5 relationship through cell-composition effects across four type-2 "
        "allergic diseases,\u201d which I would like to submit for publication as an "
        "Original Article in Allergology International.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "This study is a computational re-analysis of five publicly deposited bulk RNA-seq "
        "cohorts (GSE152004, GSE121212, GSE246323, GSE58640, GSE136825; total n = 856 "
        "patients and controls) that together span four IL5-driven allergic diseases \u2014 "
        "asthma, atopic dermatitis, eosinophilic esophagitis, and chronic rhinosinusitis "
        "with nasal polyposis. It is positioned as a hypothesis-refining, in vivo-context "
        "follow-up to our recent in vitro mechanistic report (Endo, Kurisu, Tani, Biochem "
        "Biophys Res Commun 2025), which identified the long noncoding RNA IDI2-AS1 as a "
        "negative regulator of IL5 in human T-cell-line knock-down experiments. The scope "
        "(dry re-analysis only; no new wet data) is explicitly acknowledged throughout the "
        "manuscript.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "The work contributes four findings: (i) patient-vs-control IDI2-AS1 expression is "
        "invariant at the bulk-tissue level in every cohort (random-effects pooled log\u2082FC "
        "\u2248 0; I\u00b2 = 0.26 % \u2014 a converging null across studies, not noise); "
        "(ii) sample-level IDI2-AS1 and IL5 covary positively in every testable cohort "
        "(significantly so in the largest, asthma, n = 695), which is the opposite of the "
        "within-cell repressive direction predicted by the in vitro mechanism; (iii) a "
        "bootstrap mediation-style decomposition in the asthma cohort attributes \u2248 90 % "
        "of that apparent covariation to shared expression in the eosinophil compartment, "
        "with a within-tissue direct effect indistinguishable from zero (ACME = +0.109, "
        "p = 0.002; ADE = +0.012, p = 0.76); and (iv) comparator lncRNAs and canonical "
        "type-2 cytokines in the same cohorts show the expected disease-stratified pattern, "
        "confirming that the flatness of IDI2-AS1 is a property of this specific lncRNA "
        "rather than of the pipeline.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "I believe this work is well suited for Allergology International because it "
        "directly addresses three areas central to the journal's scope: type-2 / allergic "
        "immunology (with direct relevance to the eosinophilic-asthma and CRSwNP axes on "
        "which anti-IL5 biologics act), transcriptomic methodology in human allergic "
        "disease, and the increasingly active field of non-coding RNA regulation of type-2 "
        "effector cytokines. The results also argue more generally that authentic "
        "cell-type-restricted lncRNA regulators can be systematically invisible in bulk "
        "RNA-seq catalogs of clinical tissue, and motivate prioritizing cell-type-resolved "
        "follow-up over further bulk profiling. All analysis scripts and intermediate "
        "results are publicly available on GitHub, and the study relies exclusively on "
        "public GEO datasets, enabling any reader to reproduce or extend the analysis.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "I confirm that this manuscript has not been published elsewhere and is not under "
        "consideration by another journal. I have approved the manuscript and agree with "
        "its submission to Allergology International. The author has no conflicts of "
        "interest to declare. The study involves only publicly available, fully "
        "de-identified datasets and does not require ethical approval beyond that of the "
        "original investigators.",
    )
    add_para(doc, "")

    # Correspondence block
    add_para(doc, "Please address all correspondence to:")
    add_para(doc, "")
    add_para(doc, "Hidenori Tani, Ph.D.")
    add_para(doc, "")
    add_para(doc, "Department of Health Pharmacy")
    add_para(doc, "Yokohama University of Pharmacy")
    add_para(doc, "601 Matano-cho, Totsuka-ku, Yokohama 245-0066, Japan")
    add_para(doc, "Tel.: +81 45 859 1300")
    add_para(doc, "E-mail: hidenori.tani@yok.hamayaku.ac.jp")
    add_para(doc, "ORCID: 0000-0001-6390-4136")
    add_para(doc, "")

    # Closing
    add_para(doc, "I look forward to hearing from you at your earliest convenience.")
    add_para(doc, "")
    add_para(doc, "Yours sincerely,")
    add_para(doc, "")
    add_para(doc, "Hidenori Tani, Ph.D.")

    doc.save(OUT_COVER)
    print(f"Wrote {OUT_COVER}")


# --------------------------------------------------------------------------
# 2) Tables file
# --------------------------------------------------------------------------
def _add_table(doc: Document, header: list, rows: list, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    if col_widths_cm is None:
        col_widths_cm = [2.5] * len(header)
    for i, w in enumerate(col_widths_cm):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # body rows
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)


def build_tables_doc() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- Table 1 ----
    add_para(doc,
             "Table 1. Public bulk RNA-seq cohorts re-analyzed in this study.",
             bold=True)
    t1_header = ["Disease", "GEO accession", "Tissue",
                 "Patients (n)", "Controls (n)", "Total",
                 "Platform", "Quantification", "Reference"]
    t1_rows = [
        ["Asthma", "GSE152004", "Nasal epithelium (brushing)",
         "388", "307", "695",
         "Illumina HiSeq 2000", "Raw counts", "Sajuthi et al., 2020"],
        ["Atopic dermatitis", "GSE121212",
         "Lesional skin / healthy control biopsy",
         "27", "38", "65",
         "Illumina HiSeq 2500", "Raw counts", "Tsoi et al., 2019"],
        ["Eosinophilic esophagitis", "GSE246323",
         "Esophageal biopsy (baseline)",
         "5", "5", "10",
         "Illumina NovaSeq 6000", "Raw counts", "Kleuskens et al., 2024"],
        ["Eosinophilic esophagitis", "GSE58640",
         "Esophageal biopsy",
         "10", "6", "16",
         "Illumina HiSeq 2000", "FPKM", "Sherrill et al., 2014"],
        ["Chronic rhinosinusitis with nasal polyposis", "GSE136825",
         "Nasal polyp / control inferior turbinate",
         "42", "28", "70",
         "Illumina HiSeq 4000", "Raw counts (featureCounts)",
         "Peng et al., 2019"],
        ["Combined", "", "", "472", "384", "856", "", "", ""],
    ]
    _add_table(doc, t1_header, t1_rows,
               col_widths_cm=[3.0, 2.2, 3.5, 1.4, 1.4, 1.2, 2.6, 2.4, 2.5])
    add_para(doc, "")
    add_para(
        doc,
        "Three additional candidate datasets (GSE201955, GSE65832, GSE179269) were "
        "excluded because IDI2-AS1 was absent from the available processed expression "
        "matrix or because per-sample raw data could not be recovered without "
        "re-quantification from SRA. For GSE136825, paired within-patient nasal-polyp / "
        "inferior-turbinate samples (NP_IT, n = 33) were excluded so that the case/control "
        "contrast in this cohort was structurally comparable to the four other cohorts "
        "(NP tissue vs healthy-donor inferior turbinate).",
        italic=True,
    )
    doc.add_page_break()

    # ---- Table 2 ----
    add_para(doc,
             "Table 2. Per-dataset and meta-analytic patient-vs-control differential "
             "expression for IDI2-AS1 and IL5.",
             bold=True)
    t2_header = ["Gene", "Disease", "GEO", "Method",
                 "log\u2082FC", "SE", "p\u2090\u2091"]
    t2_rows = [
        ["IDI2-AS1", "Asthma", "GSE152004",
         "DESeq2 + apeglm", "+0.0015", "0.033", "0.94"],
        ["", "Atopic dermatitis", "GSE121212",
         "DESeq2 + apeglm", "\u22120.234", "0.274", "0.41"],
        ["", "Eos. esophagitis", "GSE246323",
         "DESeq2 + apeglm", "\u22120.036", "0.145", "0.50"],
        ["", "Eos. esophagitis", "GSE58640",
         "limma-trend", "+0.247", "0.148", "0.15"],
        ["", "CRSwNP", "GSE136825",
         "DESeq2 + apeglm", "\u22122.8 \u00d7 10\u207b\u2076", "0.0014", "0.54"],
        ["", "Meta-analysis (k = 5)", "", "REML",
         "+5.1 \u00d7 10\u207b\u2075", "[\u22120.005, +0.005]", "0.98"],
        ["", "Heterogeneity", "", "", "", "I\u00b2 = 0.26 %", "Q p = 0.47"],
        ["IL5", "Asthma", "GSE152004",
         "DESeq2 + apeglm", "+0.424", "0.232", "0.014"],
        ["", "Atopic dermatitis", "GSE121212",
         "DESeq2 + apeglm", "\u22120.264", "0.325", "0.39"],
        ["", "Eos. esophagitis", "GSE58640",
         "limma-trend", "+0.609", "0.159", "0.002"],
        ["", "CRSwNP", "GSE136825",
         "DESeq2 + apeglm", "+1.0 \u00d7 10\u207b\u2076", "0.0014", "0.78"],
        ["", "Meta-analysis (k = 4)", "", "REML",
         "+0.216", "[\u22120.148, +0.580]", "0.25"],
        ["", "Heterogeneity", "", "", "",
         "I\u00b2 = 82.0 %", "Q p = 3.3 \u00d7 10\u207b\u2074"],
    ]
    _add_table(doc, t2_header, t2_rows,
               col_widths_cm=[1.8, 3.0, 2.0, 2.7, 2.0, 2.8, 2.6])
    add_para(doc, "")
    add_para(
        doc,
        "IL5 was below the rowSum \u2265 10 count threshold in GSE246323 and is therefore "
        "not included in the IL5 meta-analysis. Meta-analytic estimates use a random-effects "
        "REML model (metafor::rma).",
        italic=True,
    )
    doc.add_page_break()

    # ---- Table 3 ----
    add_para(doc,
             "Table 3. Sample-level IDI2-AS1 \u2194 IL5 association: raw, cell-type-adjusted, "
             "and mediation-style decomposition.",
             bold=True)
    t3_header = ["Cohort", "n", "Raw \u03c1", "raw p",
                 "Partial \u03c1 (| eos + Th2)", "partial p",
                 "ACME (eos.)", "ACME 95 % CI", "ACME p",
                 "ADE", "ADE p", "Prop. mediated"]
    t3_rows = [
        ["Asthma (GSE152004)", "695",
         "+0.109", "0.004",
         "+0.052", "0.17",
         "+0.109", "[+0.041, +0.176]", "0.002",
         "+0.012", "0.76", "0.90"],
        ["Eos. esophagitis (GSE58640)", "16",
         "+0.489", "0.054",
         "+0.425", "0.13",
         "\u2014", "\u2014", "\u2014",
         "\u2014", "\u2014", "\u2014"],
        ["CRSwNP (GSE136825)", "70",
         "+0.085", "0.48",
         "+0.135", "0.27",
         "\u22120.028", "[\u22120.086, +0.005]", "0.13",
         "+0.096", "0.25", "(n.s.)"],
        ["Atopic dermatitis (GSE121212)", "65",
         "+0.060", "0.63",
         "+0.010", "0.94",
         "+0.054", "[\u22120.006, +0.170]", "0.10",
         "\u22120.033", "0.78", "(n.s.)"],
    ]
    _add_table(doc, t3_header, t3_rows,
               col_widths_cm=[3.2, 0.9, 1.3, 1.1, 1.9, 1.3,
                              1.4, 2.3, 1.3, 1.2, 1.1, 1.5])
    add_para(doc, "")
    add_para(
        doc,
        "Partial Spearman correlation conditions on per-sample eosinophil and Th2 "
        "marker-signature scores (ppcor::pcor.test). The mediation-style decomposition "
        "(mediation::mediate; 1,000 nonparametric bootstrap resamples; percentile CIs) "
        "was reported only in the asthma cohort with sufficient power (n = 695); "
        "estimates from the smaller cohorts are listed for transparency but are "
        "underpowered (n < 100) and should be interpreted as directional only. Because "
        "the underlying data are cross-sectional observational measurements, the "
        "ACME / ADE values describe a statistical decomposition under the assumed "
        "mediator\u2013outcome model, not a strict counterfactual causal effect. The "
        "0.90 proportion mediated in the asthma cohort means that \u2248 90 % of the "
        "apparent IDI2-AS1 \u2194 IL5 association is statistically attributable to "
        "shared expression in the eosinophil compartment, leaving a within-tissue "
        "direct effect indistinguishable from zero (ADE p = 0.76).",
        italic=True,
    )

    doc.save(OUT_TABLES)
    print(f"Wrote {OUT_TABLES}")


# --------------------------------------------------------------------------
# 3) BBRC submission package (Cover Letter, Tables, Highlights)
# --------------------------------------------------------------------------
def build_cover_letter_bbrc() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_para(doc, "The Editor-in-Chief")
    add_para(doc, "Biochemical and Biophysical Research Communications")
    add_para(doc, "Elsevier")
    add_para(doc, "")

    add_para(doc, "24 April, 2026")
    add_para(doc, "")

    add_para(doc, "Dear Editor-in-Chief:")
    add_para(doc, "")

    add_para(
        doc,
        "Please find enclosed my manuscript, “Bulk tissue transcriptomes obscure the "
        "IDI2-AS1 / IL5 relationship through cell-composition effects across four type-2 "
        "allergic diseases,” which I would like to submit for publication as a "
        "Full-Length Research Paper in Biochemical and Biophysical Research Communications.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "This work is the direct in vivo-context companion to our recent mechanistic study "
        "in this journal — Endo, Kurisu, Tani, Biochem Biophys Res Commun 2025 — "
        "in which we identified the long noncoding RNA IDI2-AS1 as a negative regulator of "
        "interleukin 5 (IL5) through siRNA knock-down in HuT78 T-cells, with orthogonal "
        "validation in A549 epithelial cells. The natural follow-up question — whether "
        "the same regulatory axis is detectable in patient tissue from human IL5-driven "
        "allergic disease — is directly addressed here, and the two papers together "
        "form a paired in vitro / in vivo reading of one lncRNA–cytokine circuit.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "To answer this, I performed a uniform re-analysis of five public bulk RNA-seq "
        "cohorts (GSE152004, GSE121212, GSE246323, GSE58640, GSE136825; total n = 856 "
        "samples) covering four IL5-driven allergic diseases — asthma, atopic "
        "dermatitis, eosinophilic esophagitis, and chronic rhinosinusitis with nasal "
        "polyposis. The study contributes three findings: (i) patient-vs-control IDI2-AS1 "
        "expression is invariant at the bulk-tissue level in every cohort (random-effects "
        "pooled log₂FC ≈ 0; I² = 0.26 % — a converging null across "
        "studies, not noise); (ii) sample-level IDI2-AS1 and IL5 covary positively in "
        "every testable cohort (significantly so in the largest, asthma, n = 695), which "
        "is the opposite of the within-cell repressive direction predicted by the in vitro "
        "mechanism; and (iii) a bootstrap mediation-style decomposition in the asthma "
        "cohort attributes ≈ 90 % of that covariation to shared expression in the "
        "eosinophil compartment (ACME = +0.109, p = 0.002), with a within-tissue direct "
        "effect indistinguishable from zero (ADE = +0.012, p = 0.76).",
    )
    add_para(doc, "")

    add_para(
        doc,
        "Bulk patient transcriptomes therefore do not recover the predicted in vitro "
        "repressive direction. The pattern is compatible with masking of any per-cell "
        "IDI2-AS1 → IL5 effect by eosinophil-driven shifts in tissue cell "
        "composition, although we cannot exclude the alternative that the in vitro "
        "repression does not generalize to patient cells. Both possibilities require "
        "cell-type-resolved follow-up for adjudication.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "BBRC is the natural home for this paper for three reasons. First, scientific "
        "continuity: the in vitro mechanism that this study tests in vivo was published "
        "in BBRC 2025 by the same group, so readers of the 2025 paper are the natural "
        "audience. Second, methodological fit: beyond the IDI2-AS1 / IL5 case, the "
        "manuscript introduces a reusable cell-composition-adjusted re-analysis pipeline "
        "(partial correlation conditioning on marker-signature scores, followed by "
        "bootstrap mediation decomposition) directly applicable to other low-abundance, "
        "cell-type-restricted lncRNA candidates. All code is publicly released under an "
        "MIT license. Third, scope: the lncRNA–cytokine regulatory circuit, the "
        "bulk-vs-single-cell interpretive problem it exposes, and the transcriptomic "
        "decomposition framework are squarely within the biochemical / molecular biology "
        "scope of BBRC.",
    )
    add_para(doc, "")

    add_para(
        doc,
        "This is a single-author dry-analysis manuscript. The wet-lab mechanistic basis "
        "was published independently in BBRC 2025; the present study performs only "
        "computational re-analysis of publicly available datasets, generates no new "
        "sequencing data, and required no new ethics approvals beyond those of the "
        "original studies. I confirm that this manuscript has not been published "
        "elsewhere and is not under consideration by another journal. The author has no "
        "conflicts of interest to declare.",
    )
    add_para(doc, "")

    add_para(doc, "Please address all correspondence to:")
    add_para(doc, "")
    add_para(doc, "Hidenori Tani, Ph.D.")
    add_para(doc, "")
    add_para(doc, "Department of Health Pharmacy")
    add_para(doc, "Yokohama University of Pharmacy")
    add_para(doc, "601 Matano-cho, Totsuka-ku, Yokohama 245-0066, Japan")
    add_para(doc, "Tel.: +81 45 859 1300")
    add_para(doc, "E-mail: hidenori.tani@yok.hamayaku.ac.jp")
    add_para(doc, "ORCID: 0000-0001-6390-4136")
    add_para(doc, "")

    add_para(doc, "I look forward to hearing from you at your earliest convenience.")
    add_para(doc, "")
    add_para(doc, "Yours sincerely,")
    add_para(doc, "")
    add_para(doc, "Hidenori Tani, Ph.D.")

    doc.save(OUT_COVER_BBRC)
    print(f"Wrote {OUT_COVER_BBRC}")


def build_tables_bbrc() -> None:
    """BBRC tables: same content as AllergolInt, different filename."""
    import shutil
    shutil.copyfile(OUT_TABLES, OUT_TABLES_BBRC)
    print(f"Wrote {OUT_TABLES_BBRC} (copy of {OUT_TABLES.name})")


def build_highlights_bbrc() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_para(doc, "Highlights", bold=True)
    add_para(doc, "")
    # BBRC limit: ≤ 85 characters per bullet (incl. spaces). All 5 verified ≤ 85.
    bullets = [
        "Re-analysis of 5 public RNA-seq cohorts (n=856) across four IL5-driven allergies",
        "Bulk-tissue IDI2-AS1 expression is uniformly null across all cohorts (I² ≈ 0)",
        "Sample-level IDI2-AS1 covaries positively — not negatively — with IL5",
        "≈ 90 % of the covariation is mediated by the eosinophil compartment",
        "Framework is reusable for other cell-type-restricted lncRNA candidates",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.size = Pt(11)

    doc.save(OUT_HIGHLIGHTS_BBRC)
    print(f"Wrote {OUT_HIGHLIGHTS_BBRC}")


def build_declaration_of_interest_bbrc() -> None:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_para(doc, "Declaration of Interest", bold=True)
    add_para(doc, "")

    add_para(
        doc,
        "Manuscript title: Bulk tissue transcriptomes obscure the IDI2-AS1 / IL5 "
        "relationship through cell-composition effects across four type-2 allergic "
        "diseases",
    )
    add_para(doc, "")

    add_para(
        doc,
        "The author declares that he has no known competing financial interests or "
        "personal relationships that could have appeared to influence the work reported "
        "in this paper.",
    )
    add_para(doc, "")

    add_para(doc, "Hidenori Tani, Ph.D.")
    add_para(doc, "Yokohama University of Pharmacy")
    add_para(doc, "ORCID: 0000-0001-6390-4136")
    add_para(doc, "Date: 24 April, 2026")

    doc.save(OUT_DOI_BBRC)
    print(f"Wrote {OUT_DOI_BBRC}")


if __name__ == "__main__":
    build_cover_letter()
    build_tables_doc()
    build_cover_letter_bbrc()
    build_tables_bbrc()
    build_highlights_bbrc()
    build_declaration_of_interest_bbrc()
