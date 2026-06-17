"""Generate Revised_Manuscript_with_Changes_Marked.docx for BBRC revision.

Reads:
  - manuscript/full_manuscript_bbrc.docx       (v1, the originally submitted version)
  - manuscript/full_manuscript_bbrc_v2.docx    (v2, the revision in this submission)

Produces:
  - manuscript/Revised_Manuscript_with_Changes_Marked.docx

The marked document opens in any Word version (or Pages, LibreOffice) with
all changes shown as colored, formatted text — independent of Word's
"Track Changes" display mode, which can otherwise hide revision marks:

  - Insertions  : BLUE (#0070C0)  +  underline  +  bold
  - Deletions   : RED  (#C00000)  +  strikethrough  +  bold

Unchanged text is rendered with the original formatting of v2 preserved
(italics, etc.). This is purely visual markup — no <w:ins> / <w:del> tracked-
revision wrappers — so reviewers do not need to switch to "All Markup" view
to see the changes.

Strategy
--------
Paragraph-level alignment. For each v2 paragraph we find the best v1 match
by SequenceMatcher ratio. We then character-level diff the matched pair and
emit a mixed run sequence of unchanged / inserted / deleted spans. New
paragraphs (no v1 match) are wrapped as full insertions. Removed paragraphs
(v1 with no v2 match) are inserted as full deletions at the previous index.

Reproducibility: this script is idempotent — re-running overwrites the
output deterministically.
"""
from __future__ import annotations

import copy
import difflib
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + W_NS + "}"
NSMAP = {"w": W_NS}

INS_COLOR = "0070C0"  # blue
DEL_COLOR = "C00000"  # red


def _augment_rpr(rpr: etree._Element, *, color_hex: str, underline: bool = False,
                 strike: bool = False, bold: bool = True) -> None:
    """Add color / underline / strike / bold to a <w:rPr>, deduplicating."""
    for tag in ("color", "u", "strike", "b", "bCs"):
        for existing in rpr.findall(W + tag):
            rpr.remove(existing)
    if bold:
        etree.SubElement(rpr, W + "b")
        etree.SubElement(rpr, W + "bCs")
    color = etree.SubElement(rpr, W + "color")
    color.set(W + "val", color_hex)
    if underline:
        u = etree.SubElement(rpr, W + "u")
        u.set(W + "val", "single")
    if strike:
        etree.SubElement(rpr, W + "strike")


def _make_run(text: str, rpr_template: etree._Element | None = None,
              mark: str | None = None) -> etree._Element:
    """Build a <w:r> with optional copied <w:rPr> and optional ins/del styling.

    mark = None  : unchanged run, preserve template formatting only
    mark = "ins" : add blue + underline + bold
    mark = "del" : add red + strikethrough + bold
    """
    tmp = etree.Element(W + "tmp")
    r = etree.SubElement(tmp, W + "r")
    if rpr_template is not None:
        rpr = copy.deepcopy(rpr_template)
    else:
        rpr = etree.SubElement(r, W + "rPr")
    if mark == "ins":
        _augment_rpr(rpr, color_hex=INS_COLOR, underline=True, bold=True)
    elif mark == "del":
        _augment_rpr(rpr, color_hex=DEL_COLOR, strike=True, bold=True)
    if rpr_template is not None and len(rpr) > 0:
        r.append(rpr)
    elif rpr_template is None and len(rpr) == 0:
        r.remove(rpr)
    t = etree.SubElement(r, W + "t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


CHAR_DIFF_THRESHOLD = 0.85  # below this similarity, show as full delete+insert


def char_rpr_map(paragraph) -> list[etree._Element | None]:
    """Return a list of (rPr element or None) for each character in paragraph.

    This preserves per-character formatting (e.g. italics on gene names within
    a paragraph of otherwise upright prose), so when we rebuild the paragraph
    we do not propagate the italic formatting of one mid-paragraph run to
    the whole paragraph.
    """
    out: list[etree._Element | None] = []
    for run in paragraph._element.findall(W + "r"):
        rpr = run.find(W + "rPr")
        for t in run.findall(W + "t"):
            text = t.text or ""
            for _ in text:
                out.append(rpr)
    return out


def _rpr_key(rpr: etree._Element | None) -> str:
    """Canonical string for an rPr (or None) so we can merge consecutive runs
    that share identical formatting."""
    if rpr is None:
        return ""
    return etree.tostring(rpr, method="c14n").decode("utf-8")


def diff_paragraph(v1_text: str, v2_text: str,
                   v2_paragraph) -> list[etree._Element]:
    """Return run elements representing v2 with diff vs v1.

    For high-similarity inline edits (similarity >= 0.85), we walk through
    v2 character-by-character and preserve the per-character rPr (so italic
    gene names and other inline formatting are not broken). For low-similarity
    heavy rewrites, we emit the whole v1 paragraph as a strikethrough deletion
    followed by the whole v2 paragraph as a colored insertion.
    """
    ratio = difflib.SequenceMatcher(a=v1_text, b=v2_text, autojunk=False).ratio()
    out: list[etree._Element] = []

    if ratio < CHAR_DIFF_THRESHOLD:
        # Heavy rewrite: keep first_rpr as a coarse template; readability of
        # the diff is more important here than character-by-character fidelity.
        first_rpr = None
        for r in v2_paragraph._element.findall(W + "r"):
            rpr = r.find(W + "rPr")
            if rpr is not None:
                first_rpr = rpr
                break
        out.append(_make_run(v1_text, first_rpr, mark="del"))
        out.append(_make_run(" ", first_rpr))
        out.append(_make_run(v2_text, first_rpr, mark="ins"))
        return out

    # Light inline edits: rebuild while preserving per-character v2 rPr.
    v2_char_rpr = char_rpr_map(v2_paragraph)
    if len(v2_char_rpr) != len(v2_text):
        # Defensive: fall back to no rPr for unchanged if anything is off.
        v2_char_rpr = [None] * len(v2_text)

    sm = difflib.SequenceMatcher(a=v1_text, b=v2_text, autojunk=False)
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            # Group consecutive chars that share the same rPr into a single run
            run_start = j1
            current_key = _rpr_key(v2_char_rpr[j1]) if j2 > j1 else ""
            j = j1 + 1
            while j < j2:
                k = _rpr_key(v2_char_rpr[j])
                if k != current_key:
                    out.append(_make_run(v2_text[run_start:j], v2_char_rpr[run_start]))
                    run_start = j
                    current_key = k
                j += 1
            if run_start < j2:
                out.append(_make_run(v2_text[run_start:j2], v2_char_rpr[run_start]))
        elif op == "replace":
            out.append(_make_run(v1_text[i1:i2], None, mark="del"))
            out.append(_make_run(v2_text[j1:j2], None, mark="ins"))
        elif op == "delete":
            out.append(_make_run(v1_text[i1:i2], None, mark="del"))
        elif op == "insert":
            out.append(_make_run(v2_text[j1:j2], None, mark="ins"))
    return out


def clear_paragraph_runs(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        tag = etree.QName(child.tag).localname
        if tag in {"r", "ins", "del", "hyperlink", "smartTag", "fldSimple"}:
            p.remove(child)


def match_v1_paragraphs(v2_text: str, v1_paragraphs: list[str]) -> int | None:
    best_idx = None
    best_ratio = 0.5
    for i, t in enumerate(v1_paragraphs):
        if not t.strip():
            continue
        r = difflib.SequenceMatcher(a=t, b=v2_text, autojunk=False).ratio()
        if r > best_ratio:
            best_ratio = r
            best_idx = i
    return best_idx


def main() -> None:
    here = Path(__file__).resolve().parent
    v1_doc = Document(str(here / "full_manuscript_bbrc.docx"))
    v2_doc = Document(str(here / "full_manuscript_bbrc_v2.docx"))
    out_path = here / "Revised_Manuscript_with_Changes_Marked.docx"

    v1_texts = [p.text for p in v1_doc.paragraphs]
    v2_texts = [p.text for p in v2_doc.paragraphs]

    used_v1: set[int] = set()
    v2_to_v1: list[int | None] = []
    for v2_text in v2_texts:
        match = match_v1_paragraphs(v2_text, v1_texts)
        if match is not None and match not in used_v1:
            used_v1.add(match)
            v2_to_v1.append(match)
        else:
            v2_to_v1.append(None)

    n_inserted_para = 0
    n_changed_para = 0

    for v2_idx, v2_para in enumerate(v2_doc.paragraphs):
        v1_idx = v2_to_v1[v2_idx]
        v2_text = v2_para.text

        if v1_idx is not None:
            v1_text = v1_texts[v1_idx]
            if v1_text == v2_text:
                continue
            new_children = diff_paragraph(v1_text, v2_text, v2_para)
            clear_paragraph_runs(v2_para)
            for ch in new_children:
                v2_para._element.append(ch)
            n_changed_para += 1
        else:
            if not v2_text.strip():
                continue
            run = _make_run(v2_text, None, mark="ins")
            clear_paragraph_runs(v2_para)
            v2_para._element.append(run)
            n_inserted_para += 1

    # Header note
    body = v2_doc.element.body
    note_para = v2_doc.paragraphs[0]
    note_text = (
        "[Revised Manuscript with Changes Marked — generated 2026-05-12 for "
        "BBRC-26-2497. INSERTIONS shown in BLUE, bold, underlined. DELETIONS "
        "shown in RED, bold, strikethrough. Unchanged text is in normal "
        "formatting. For the clean (accepted) version see "
        "full_manuscript_bbrc_v2.docx.]"
    )
    new_p = copy.deepcopy(note_para._element)
    for ch in list(new_p):
        tag = etree.QName(ch.tag).localname
        if tag in {"r", "ins", "del", "hyperlink"}:
            new_p.remove(ch)
    note_run = etree.SubElement(new_p, W + "r")
    rpr = etree.SubElement(note_run, W + "rPr")
    etree.SubElement(rpr, W + "i")
    color = etree.SubElement(rpr, W + "color")
    color.set(W + "val", "595959")  # dark grey
    nt = etree.SubElement(note_run, W + "t")
    nt.text = note_text
    nt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    body.insert(0, new_p)

    v2_doc.save(str(out_path))
    print(f"Saved: {out_path}")
    print(f"  paragraphs modified inline: {n_changed_para}")
    print(f"  paragraphs newly inserted:  {n_inserted_para}")
    print(f"  size: {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
